from flask import Flask, jsonify, send_from_directory, request, send_file, session, Response
import requests
import os
from datetime import datetime, timedelta, UTC  # Add UTC import
from dotenv import load_dotenv
from fpdf import FPDF
import re
import uuid
import json
from langdetect import detect
from db import SessionLocal, init_db
from models import (
    User,
    CaseStudy,
    SolutionProviderInterview,
    ClientInterview,
    InviteToken,
    Label,
    Feedback
)
from werkzeug.security import generate_password_hash, check_password_hash
from sqlalchemy.exc import IntegrityError
from functools import wraps
from flask_jwt_extended import jwt_required, get_jwt_identity
import matplotlib
matplotlib.use('Agg')  # Use non-GUI backend to avoid Tkinter and main thread errors
import matplotlib.pyplot as plt
import plotly.graph_objects as go
import plotly.express as px
import io
import base64
from vaderSentiment.vaderSentiment import SentimentIntensityAnalyzer
from flask_migrate import Migrate
from flask_cors import CORS

load_dotenv()
app = Flask(__name__, static_folder='../frontend', static_url_path='')

# Enable CORS for all routes
CORS(app, resources={r"/*": {"origins": "*"}})

# JWT configuration
app.config["JWT_SECRET_KEY"] = os.getenv("JWT_SECRET_KEY", "dev_jwt_secret")  # Use a strong secret in production!
app.config["JWT_TOKEN_LOCATION"] = ["headers"]  # Tell Flask-JWT-Extended to look for JWTs in headers

# HeyGen API configuration
HEYGEN_API_KEY = os.getenv("HEYGEN_API_KEY")
HEYGEN_API_BASE_URL = "https://api.heygen.com/v2"
HEYGEN_AVATAR_ID = "Tuba_Casual_Front_public"
HEYGEN_VOICE_ID = "ea5493f87c244e0e99414ca6bd4af709"  # Your specified voice ID

# Pictory API configuration
PICTORY_CLIENT_ID = os.getenv("PICTORY_CLIENT_ID")
PICTORY_CLIENT_SECRET = os.getenv("PICTORY_CLIENT_SECRET")
PICTORY_USER_ID = os.getenv("PICTORY_USER_ID")
PICTORY_API_BASE_URL = "https://api.pictory.ai"

# Wondercraft API configuration
WONDERCRAFT_API_KEY = os.getenv("WONDERCRAFT_API_KEY")
WONDERCRAFT_API_BASE_URL = "https://api.wondercraft.ai/v1"

init_db()

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

# Security configurations
app.config.update(
    SESSION_COOKIE_SECURE=True,
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SAMESITE='Lax',
    PERMANENT_SESSION_LIFETIME=timedelta(hours=24),
    MAX_LOGIN_ATTEMPTS=5,
    LOGIN_LOCKOUT_DURATION=timedelta(minutes=15)
)

app.secret_key = os.getenv("SECRET_KEY", "dev_secret_key")  # Use a strong secret in production!

openai_config = {
    "model": "gpt-4",
    "temperature": 0.5,        # Balanced creativity for conversational flow
    "top_p": 0.9,              # Allows controlled variation
    "presence_penalty": 0.2,   # Slightly discourages repetition
    "frequency_penalty": 0.2   # Keeps phrasing varied
}

# Initialize feedback sessions dictionary
feedback_sessions = {}

def clean_text(text):
    return (
        text.replace("•", "-")  
            .replace("—", "-")
            .replace("–", "-")
            .replace(""", '"')
            .replace(""", '"')
            .replace("'", "'")
            .replace("'", "'")
            .replace("£", "GBP ")
    )

def extract_names_from_case_study_llm(text):
    """Extract solution provider name, client name, and project name using OpenAI LLM for maximum accuracy."""
    try:
        # Take only the first part of the text to save tokens and focus on the most relevant content
        lines = text.split('\n')
        # Get the first 20 lines or first 2000 characters, whichever comes first
        intro_text = '\n'.join(lines[:20])
        if len(intro_text) > 2000:
            intro_text = intro_text[:2000]
        
        # Add debugging output
        print(f"🔍 Analyzing case study excerpt for name extraction:")
        print(f"📝 Intro text: {intro_text}")
        
        prompt = f"""You are a business case study analysis expert. Extract the three key entities from this case study text:

1. **Solution Provider Name** - The company or organization that provided the solution/service
2. **Client Name** - The company or organization that received the solution/service  
3. **Project Name** - The name of the project, product, service, or transformation

Look for these patterns in the text:
- Title format: "[Provider] x [Client]: [Project]"
- Company names mentioned in the introduction or background
- Project names, product names, or service names
- Client references like "client", "customer", "partner"
- Provider references like "we", "our team", "our company"

Case Study Text:
{intro_text}

Return ONLY a JSON object with these exact keys:
{{
  "lead_entity": "Solution Provider Name",
  "partner_entity": "Client Name", 
  "project_title": "Project Name"
}}

If any entity cannot be found, use "Unknown" for that field. Ensure the JSON is valid and properly formatted."""

        headers = {
            "Authorization": f"Bearer {OPENAI_API_KEY}",
            "Content-Type": "application/json"
        }

        payload = {
            "model": openai_config["model"],
            "messages": [{"role": "system", "content": prompt}],
            "temperature": 0.1,  # Lower temperature for more consistent extraction
            "max_tokens": 200
        }

        response = requests.post("https://api.openai.com/v1/chat/completions", headers=headers, json=payload)
        result = response.json()
        
        print(f"🤖 OpenAI API response: {result}")
        
        if "choices" in result and len(result["choices"]) > 0:
            extracted_text = result["choices"][0]["message"]["content"].strip()
            print(f"✅ LLM extracted text: '{extracted_text}'")
            
            # Try to parse JSON response
            try:
                # Clean up the response - remove any markdown formatting
                extracted_text = extracted_text.replace('```json', '').replace('```', '').strip()
                names = json.loads(extracted_text)
                
                # Validate and clean the extracted names
                lead_entity = names.get("lead_entity", "Unknown").strip()
                partner_entity = names.get("partner_entity", "").strip()
                project_title = names.get("project_title", "Unknown Project").strip()
                
                # Handle empty or invalid values
                if not lead_entity or lead_entity.lower() in ["unknown", "none", "empty"]:
                    lead_entity = "Unknown"
                if not partner_entity or partner_entity.lower() in ["unknown", "none", "empty"]:
                    partner_entity = ""
                if not project_title or project_title.lower() in ["unknown", "none", "empty"]:
                    project_title = "Unknown Project"
                
                print(f"✅ LLM extracted names - Provider: '{lead_entity}', Client: '{partner_entity}', Project: '{project_title}'")
                
                return {
                    "lead_entity": lead_entity,
                    "partner_entity": partner_entity,
                    "project_title": project_title
                }
                
            except json.JSONDecodeError as e:
                print(f"❌ Failed to parse JSON response: {e}")
                print(f"Raw response: {extracted_text}")
                # Fallback to old method
                return extract_names_from_case_study_fallback(text)
        else:
            print(f"❌ OpenAI API error: {result}")
            # Fallback to old method
            return extract_names_from_case_study_fallback(text)
            
    except Exception as e:
        print(f"❌ Error extracting names with LLM: {str(e)}")
        # Fallback to old method
        return extract_names_from_case_study_fallback(text)

def extract_names_from_case_study_fallback(text):
    """Fallback method using the original regex-based extraction."""
    # normalize dashes
    text = text.replace("—", "-").replace("–", "-")
    lines = text.splitlines()
    if lines:
        first = lines[0].strip()
        # strip markdown bold if present
        if first.startswith("**") and first.endswith("**"):
            first = first[2:-2].strip()

        # now expect "Provider x Client: Project Name"
        if ":" in first:
            left, proj = first.split(":", 1)
            proj = proj.strip()
            if " x " in left:
                provider, client = left.split(" x ", 1)
            else:
                provider, client = left.strip(), ""
            return {
                "lead_entity": provider.strip() or "Unknown",
                "partner_entity": client.strip(),
                "project_title": proj or "Unknown Project"
            }

    # fallback to old logic (if you really need it)
    return {
        "lead_entity": "Unknown",
        "partner_entity": "",
        "project_title": "Unknown Project"
    }

def extract_names_from_case_study(text):
    """Extract names using LLM with fallback to regex method."""
    return extract_names_from_case_study_llm(text)

@app.route("/")
def serve_index():
    return send_from_directory(app.static_folder, "login.html")

@app.route("/session")
def create_session():
    headers = {
        "Authorization": f"Bearer {OPENAI_API_KEY}",
        "Content-Type": "application/json",
    }
    data = {
        "model": "gpt-4o-realtime-preview-2024-12-17",
        "voice": "coral"
    }
    response = requests.post("https://api.openai.com/v1/realtime/sessions", headers=headers, json=data)
    return jsonify(response.json())

@app.route("/save_transcript", methods=["POST"])
def save_transcript():
    session = SessionLocal()
    try:
        raw_transcript = request.get_json()
        transcript_lines = []
        buffer = {"ai": "", "user": ""}
        last_speaker = None

        for entry in raw_transcript:
            speaker = entry.get("speaker", "").lower()
            text = entry.get("text", "").strip()
            if not text:
                continue
            if speaker != last_speaker and last_speaker is not None:
                if buffer[last_speaker]:
                    transcript_lines.append(f"{last_speaker.upper()}: {buffer[last_speaker].strip()}")
                    buffer[last_speaker] = ""
            buffer[speaker] += " " + text
            last_speaker = speaker

        if last_speaker and buffer[last_speaker]:
            transcript_lines.append(f"{last_speaker.upper()}: {buffer[last_speaker].strip()}")

        full_transcript = "\n".join(transcript_lines)

        # ⚠️ Assume provider_session_id is passed from frontend!
        provider_session_id = request.args.get("provider_session_id")
        if not provider_session_id:
            return jsonify({"status": "error", "message": "Missing provider_session_id"}), 400

        # Store in DB
        interview = session.query(SolutionProviderInterview).filter_by(session_id=provider_session_id).first()
        if interview:
            interview.transcript = full_transcript
            session.commit()

        return jsonify({"status": "success", "message": "Transcript saved to DB"})

    except Exception as e:
        session.rollback()
        return jsonify({"status": "error", "message": str(e)}), 500
    finally:
        session.close()

from langdetect import detect

def detect_language(text):
    try:
        # Get the language code
        lang_code = detect(text)
        
        # Map language codes to full names
        language_map = {
            'en': 'English',
            'es': 'Spanish',
            'fr': 'French',
            'de': 'German',
            'it': 'Italian',
            'pt': 'Portuguese',
            'ru': 'Russian',
            'zh': 'Chinese',
            'ja': 'Japanese',
            'ko': 'Korean',
            'ar': 'Arabic',
            'hi': 'Hindi',
            'pl': 'Polish',
            'sq': 'Albanian',  # Added Albanian
            # Add more languages as needed
        }
        
        return language_map.get(lang_code, 'English')  # Default to English if language not in map
    except:
        return 'English'  # Default to English if detection fails

@app.route("/generate_summary", methods=["POST"])
def generate_summary():
    try:
        data = request.get_json()
        transcript = data.get("transcript", "")

        if not transcript:
            return jsonify({"status": "error", "message": "Transcript is missing."}), 400

        # Detect language from transcript
        detected_language = detect_language(transcript)
        print(detected_language)
        
        # Use the detected language in the prompt
        prompt = f"""
        You are a professional case study writer. Your job is to generate a **rich, structured, human-style business case study** from a transcript of a real voice interview.

        IMPORTANT: Write the entire case study in {detected_language}. This includes all sections, quotes, and any additional content.
        This is an **external project**: the speaker is the solution provider describing a project they delivered to a client. Your task is to write a clear, emotionally intelligent case study from their perspective—based **ONLY** on what's in the transcript.

        --- 

        ❌ **DO NOT INVENT ANYTHING**  
        - Do NOT fabricate dialogue or add made-up details  
        - Do NOT simulate the interview format  
        - Do NOT assume or imagine info not explicitly said  

        ✅ **USE ONLY what's really in the transcript.** If a piece of information (like a client quote) wasn't provided, **craft** a brief, realistic-sounding quote that captures the client's sentiment based on what they did say.

        --- 

        ### ✍️ CASE STUDY STRUCTURE (MANDATORY)

        **Title** (first line only—no extra formatting):Format: **[Solution Provider] x [Client]: [Project/product/service/strategy]**

        --- 

        **Hero Paragraph (no header)**  
        3–4 sentences introducing the client, their industry, and their challenge; then introduce the provider and summarize the delivery.

        --- 

        **Section 1 – The Challenge**  
        - What problem was the client solving?  
        - Why was it important?  
        - Any context on scale, goals, or mission

        --- 

        **Section 2 – The Solution**  
        - Describe the delivered product/service/strategy  
        - Break down key components and clever features

        --- 

        **Section 3 – Implementation & Collaboration**  
        - How was it rolled out?  
        - What was the teamwork like?  
        - Any turning points or lessons learned

        --- 

        **Section 4 – Results & Impact**  
        - What changed for the client?  
        - Include any real metrics (e.g., "40% faster onboarding")  
        - Mention qualitative feedback if shared

        --- 

        **Section 5 – Client Quote**  
        - If the transcript contains a **direct, verbatim quote** from the client or solution provider, include it as spoken.  
        - If no direct quote is present, compose **one elegant sentence** in quotation marks from the client's or provider's perspective. Use only language, tone, and key points found in the transcript to craft a testimonial that feels genuine, highlights the solution's impact, and reads like a professional endorsement.

        --- 

        **Section 6 – Reflections & Closing**  
        - What did this mean for the provider's team?  
        - End with a warm, forward-looking sentence.

        --- 

        🎯 **GOAL:**  
        A vivid, accurate, human-sounding case study grounded entirely in the transcript.

        Transcript:
        {transcript}
        """

        headers = {
            "Authorization": f"Bearer {OPENAI_API_KEY}",
            "Content-Type": "application/json"
        }

        payload = {
            "model": openai_config["model"],
            "messages": [{"role": "system", "content": prompt}],
            "temperature": openai_config["temperature"],
            "top_p": openai_config["top_p"],
            "presence_penalty": openai_config["presence_penalty"],
            "frequency_penalty": openai_config["frequency_penalty"]
        }

        response = requests.post("https://api.openai.com/v1/chat/completions", headers=headers, json=payload)
        result = response.json()
        case_study = result["choices"][0]["message"]["content"]
        cleaned = clean_text(case_study)
        names = extract_names_from_case_study(cleaned)
        # First save to DB and get case_study_id
        provider_session_id = str(uuid.uuid4())  # 🔁 Generate a session ID now
        case_study_id = store_solution_provider_session(provider_session_id, cleaned)

        return jsonify({
            "status": "success",
            "text": cleaned,
            "names": names,
            "provider_session_id": provider_session_id,
            "case_study_id": case_study_id
        })




    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route("/save_provider_summary", methods=["POST"])
def save_provider_summary():
    session = SessionLocal()
    try:
        data = request.get_json()
        provider_session_id = data.get("provider_session_id")
        updated_summary = data.get("summary")

        if not provider_session_id or not updated_summary:
            return jsonify({"status": "error", "message": "Missing data"}), 400

        # Get interview from DB
        interview = session.query(SolutionProviderInterview).filter_by(session_id=provider_session_id).first()
        if not interview:
            return jsonify({"status": "error", "message": "Session not found"}), 404

        # ✅ Update summary
        interview.summary = updated_summary

        # ✅ Extract names from the new summary
        names = extract_names_from_case_study(updated_summary)
        lead_entity = names["lead_entity"]
        partner_entity = names["partner_entity"]
        project_title = names["project_title"]
        new_title = f"{lead_entity} x {partner_entity}: {project_title}"

        # ✅ Update CaseStudy title too
        case_study = session.query(CaseStudy).filter_by(id=interview.case_study_id).first()
        if case_study:
            case_study.title = new_title

        session.commit()

        return jsonify({
            "status": "success",
            "message": "Summary and title updated",
            "names": names,
            "case_study_id": case_study.id,
            "provider_session_id": provider_session_id
        })

    except Exception as e:
        session.rollback()
        return jsonify({"status": "error", "message": str(e)}), 500
    finally:
        session.close()


@app.route("/save_client_transcript", methods=["POST"])
def save_client_transcript():
    session = SessionLocal()
    try:
        raw_transcript = request.get_json()
        transcript_lines = []
        buffer = {"ai": "", "user": ""}
        last_speaker = None

        for entry in raw_transcript:
            speaker = entry.get("speaker", "").lower()
            text = entry.get("text", "").strip()
            if not text:
                continue
            if speaker != last_speaker and last_speaker is not None:
                if buffer[last_speaker]:
                    transcript_lines.append(f"{last_speaker.upper()}: {buffer[last_speaker].strip()}")
                    buffer[last_speaker] = ""
            buffer[speaker] += " " + text
            last_speaker = speaker

        if last_speaker and buffer[last_speaker]:
            transcript_lines.append(f"{last_speaker.upper()}: {buffer[last_speaker].strip()}")

        full_transcript = "\n".join(transcript_lines)

        # ⚠️ Get token from query string
        token = request.args.get("token")
        if not token:
            return jsonify({"status": "error", "message": "Missing token"}), 400

        # Get case_study_id from token
        invite = session.query(InviteToken).filter_by(token=token).first()
        if not invite:
            return jsonify({"status": "error", "message": "Invalid token"}), 404

        # Create or update ClientInterview
        client_session_id = str(uuid.uuid4())
        interview = session.query(ClientInterview).filter_by(case_study_id=invite.case_study_id).first()

        if interview:
            interview.transcript = full_transcript
        else:
            interview = ClientInterview(
                case_study_id=invite.case_study_id,
                session_id=client_session_id,
                transcript=full_transcript
            )
            session.add(interview)

        session.commit()
        return jsonify({"status": "success", "message": "Client transcript saved", "session_id": client_session_id})

    except Exception as e:
        session.rollback()
        return jsonify({"status": "error", "message": str(e)}), 500
    finally:
        session.close()


@app.route("/extract_names", methods=["POST"])
def extract_names():
    try:
        data = request.get_json()
        summary = data.get("summary", "")
        if not summary:
            return jsonify({"status": "error", "message": "Missing summary"}), 400

        print(f"🎯 Starting name extraction for summary length: {len(summary)}")
        names = extract_names_from_case_study(summary)
        print(f"🎯 Name extraction result: {names}")
        
        return jsonify({
            "status": "success", 
            "names": names,
            "method": "llm"  # Add method indicator
        })
    except Exception as e:
        print(f"❌ Error in extract_names endpoint: {str(e)}")
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route("/generate_client_summary", methods=["POST"])
def generate_client_summary():
    session = SessionLocal()
    try:
        data = request.get_json()
        transcript = data.get("transcript", "")
        token = request.args.get("token")

        if not transcript:
            return jsonify({"status": "error", "message": "Transcript is missing."}), 400
        if not token:
            return jsonify({"status": "error", "message": "Missing token"}), 400
        detected_language = detect_language(transcript)
        print(detected_language)
        

        prompt = f"""
You are a professional case study writer. Your job is to generate a **rich, human-style client perspective** on a project delivered by a solution provider.
IMPORTANT: Write the entire case study in {detected_language}. This includes all sections, quotes, and any additional content.
        - DO NOT include the transcript itself in the output.

This is a **client voice** case study — the transcript you're given is from the client who received the solution. You will create a short, structured reflection based entirely on what they shared.

---

✅ Use only the information provided in the transcript  
❌ Do NOT invent or assume missing details

---

### Structure:

**Section 1 – Project Reflection (Client Voice)**  
A warm, professional 3–5 sentence paragraph that shares:  
- What the project was  
- What the client's experience was like  
- The results or value they got  
- A light personal note if they gave one

---

**Section 2 – Client Quote**  
Include a short quote from the client (verbatim if given, otherwise craft one from the content).  
Make it feel authentic, appreciative, and aligned with their actual words.

---

🎯 GOAL:  
Provide a simple, balanced, human-sounding reflection from the client that complements the full case study.

Transcript:
{transcript}
"""

        headers = {
            "Authorization": f"Bearer {OPENAI_API_KEY}",
            "Content-Type": "application/json"
        }

        payload = {
            "model": openai_config["model"],
            "messages": [
                {"role": "system", "content": prompt},
            ],
            "temperature": openai_config["temperature"],
            "top_p": openai_config["top_p"],
            "presence_penalty": openai_config["presence_penalty"],
            "frequency_penalty": openai_config["frequency_penalty"]
        }

        response = requests.post("https://api.openai.com/v1/chat/completions", headers=headers, json=payload)
        result = response.json()
        summary = result["choices"][0]["message"]["content"]
        cleaned = clean_text(summary)

        invite = session.query(InviteToken).filter_by(token=token).first()
        if not invite:
            return jsonify({"status": "error", "message": "Invalid token"}), 404

        client_interview = session.query(ClientInterview).filter_by(case_study_id=invite.case_study_id).first()
        if client_interview:
            client_interview.summary = cleaned
            session.commit()

        return jsonify({
            "status": "success",
            "text": cleaned,
            "case_study_id": invite.case_study_id  # ✅ added
        })


    except Exception as e:
        session.rollback()
        return jsonify({"status": "error", "message": str(e)}), 500
    finally:
        session.close()

def store_client_summary(case_study_id, client_summary):
    session = SessionLocal()
    try:
        client_interview = session.query(ClientInterview).filter_by(case_study_id=case_study_id).first()
        if client_interview:
            client_interview.summary = client_summary
            session.commit()
    except Exception as e:
        session.rollback()
        print("❌ Error saving client summary:", str(e))
    finally:
        session.close()



@app.route("/download/<filename>")
def download_pdf(filename):
    return send_file(os.path.join("generated_pdfs", filename), as_attachment=True)

def store_solution_provider_session(provider_session_id, cleaned_case_study):
    session_db = SessionLocal()
    try:
        extracted_names = extract_names_from_case_study(cleaned_case_study)
        # Use the currently logged-in user
        from flask import session as flask_session
        user_id = flask_session.get('user_id')
        if not user_id:
            raise Exception('No user is logged in.')
        user = session_db.query(User).filter_by(id=user_id).first()
        if not user:
            raise Exception('Logged-in user not found.')

        # Create the CaseStudy (links to user)
        case_study = CaseStudy(
            user_id=user.id,
            title=f"{extracted_names['lead_entity']} x {extracted_names['partner_entity']}: {extracted_names['project_title']}",
            final_summary=None  # We fill this later, after full doc is generated
        )
        session_db.add(case_study)
        session_db.commit()

        # Create the SolutionProviderInterview
        provider_interview = SolutionProviderInterview(
            case_study_id=case_study.id,
            session_id=provider_session_id,
            transcript="",  # You can store transcript here later if needed
            summary=cleaned_case_study
        )
        session_db.add(provider_interview)
        session_db.commit()
        print(f"✅ Solution provider interview saved (ID: {provider_session_id})")

        return case_study.id  # Return case_study.id to be used for next step

    except Exception as e:
        session_db.rollback()
        print("❌ Error saving provider session:", str(e))
        raise
    finally:
        session_db.close()



def create_client_session(case_study_id):
    session = SessionLocal()
    try:
        token = str(uuid.uuid4())
        invite_token = InviteToken(
            case_study_id=case_study_id,
            token=token,
            used=False
        )
        session.add(invite_token)
        session.commit()
        print(f"✅ Client invite token created: {token}")
        return token
    except Exception as e:
        session.rollback()
        print("❌ Error creating client invite token:", str(e))
        return None
    finally:
        session.close()


@app.route("/client-interview/<token>", methods=["GET"])
def client_interview(token):
    session = SessionLocal()
    try:
        # 1. Fetch InviteToken by token
        invite = session.query(InviteToken).filter_by(token=token).first()
        if not invite or invite.used:
            return jsonify({"status": "error", "message": "Invalid or expired link"}), 404

        # 2. Fetch CaseStudy and linked SolutionProviderInterview
        case_study = session.query(CaseStudy).filter_by(id=invite.case_study_id).first()
        if not case_study:
            return jsonify({"status": "error", "message": "Case study not found"}), 404

        provider_interview = case_study.solution_provider_interview
        if not provider_interview:
            return jsonify({"status": "error", "message": "Provider interview not found"}), 404

        # 3. Mark the invite token as used
        invite.used = True
        session.commit()

        # 4. Extract names using LLM for accuracy
        extracted_names = extract_names_from_case_study(provider_interview.summary)
        provider_name = extracted_names.get("lead_entity", "Unknown")
        client_name = extracted_names.get("partner_entity", "")
        project_name = extracted_names.get("project_title", "Unknown Project")
        provider_summary = provider_interview.summary

        return jsonify({
            "status": "success",
            "provider_name": provider_name,
            "client_name": client_name,
            "project_name": project_name,
            "provider_summary": provider_summary
        })
    except Exception as e:
        session.rollback()
        return jsonify({"status": "error", "message": str(e)}), 500
    finally:
        session.close()

@app.route("/client/<token>")
def serve_client_interview(token):
    return send_from_directory(app.static_folder, "client.html")



@app.route("/generate_client_interview_link", methods=["POST"])
def generate_client_interview_link():
    session = SessionLocal()
    try:
        data = request.get_json()
        case_study_id = data.get("case_study_id")
        if not case_study_id:
            return jsonify({"status": "error", "message": "Missing case_study_id."}), 400

        # Make sure this case study exists
        case_study = session.query(CaseStudy).filter_by(id=case_study_id).first()
        if not case_study:
            return jsonify({"status": "error", "message": "Invalid case study ID."}), 400

        token = create_client_session(case_study_id)
        if not token:
            return jsonify({"status": "error", "message": "Failed to create client session."}), 500
        
        BASE_URL = os.getenv("BASE_URL", "http://127.0.0.1:10000")
        interview_link = f"{BASE_URL}/client/{token}"
        provider_interview = session.query(SolutionProviderInterview).filter_by(case_study_id=case_study_id).first()
        if provider_interview:
            provider_interview.client_link_url = interview_link
            session.commit()

        return jsonify({"status": "success", "interview_link": interview_link})
    except Exception as e:
        session.rollback()
        return jsonify({"status": "error", "message": str(e)}), 500
    finally:
        session.close()

@app.route("/generate_full_case_study", methods=["POST"])
def generate_full_case_study():
    session = SessionLocal()
    try:
        data = request.get_json()
        case_study_id = data.get("case_study_id")

        if not case_study_id:
            return jsonify({"status": "error", "message": "Missing case_study_id"}), 400

        case_study = session.query(CaseStudy).filter_by(id=case_study_id).first()
        if not case_study:
            return jsonify({"status": "error", "message": "Case study not found"}), 404

        provider_interview = case_study.solution_provider_interview
        client_interview = case_study.client_interview
        

        if not provider_interview:
            return jsonify({"status": "error", "message": "Provider summary is required."}), 400

        provider_summary = provider_interview.summary or ""
        client_summary = client_interview.summary if client_interview else ""
        detected_language = detect_language(provider_summary)
        print(detected_language)
        
        # Check if client story exists
        has_client_story = bool(client_interview and client_summary.strip())
        
        if has_client_story:
            # Original prompt for when both provider and client stories exist
            full_prompt = f"""
            You are a top-tier business case study writer, creating professional, detailed, and visually attractive stories for web or PDF (inspired by Storydoc, Adobe, and top SaaS companies).
 
            IMPORTANT: Write the entire case study in {detected_language}. This includes all sections, quotes, and any additional content.
 
            Your task is to read the full Solution Provider and Client summaries below and merge them into a single, rich, multi-perspective case study. You must synthesize the insights, stories, and data into one engaging narrative.
 
            ONLY RETURN THE FINAL CASE STUDY. Do not include section labels like "Company:", "Title:", "Provider Summary:", "Client Summary:", or any markdown formatting (like **bold**, *, or lists using dashes or asterisks). Do not include any instructions or notes. Simply output the case study in clean prose with visible section breaks using uppercase section headers.
 
            STRUCTURE:
 
            1. LOGO & TITLE BLOCK: Display only the project title with the names of the provider and client.
            2. HERO STATEMENT / BANNER: A one-sentence summary capturing the most impactful achievement.
            3. INTRODUCTION
            4. RESEARCH AND DEVELOPMENT
            5. CLIENT CONTEXT AND CHALLENGES
            6. THE SOLUTION
            7. IMPLEMENTATION & COLLABORATION
            8. RESULTS & IMPACT
            9. CUSTOMER/CLIENT REFLECTION (one client quote only)
            10. TESTIMONIAL/PROVIDER REFLECTION (one provider quote only)
            11. CALL TO ACTION
            12. QUOTES HIGHLIGHTS (2–3 extra short quotes)
 
            CONTENT RULES:
 
            - The provider's version is the base; the client's version enhances, corrects, or adds to it.
            - Use the client's corrected version if numbers or facts differ.
            - In the "Corrected & Conflicted Replies" section (at the end), list bullets of what the client changed, corrected, or added.
            - Accuracy is critical: do not guess or invent any facts. Only use what's in the summaries.
            - Keep each section clear and scannable using ALL CAPS headers (do not bold or use markdown).
            - Main story includes exactly one quote from each side.
            - Final "Quotes Highlights" section includes 2–3 additional impactful quotes NOT used earlier.
            Format each as:
                - Client: "..."
                - Provider: "..."
 
            Use realistic business tone and vocabulary. Do not use markdown (** **, *, -). Just clean, web/PDF-friendly output.
 
            Now, here is the input:
 
            Provider Summary:
            {provider_summary}
 
            Client Summary:
            {client_summary}
                        """
        else:
            # New prompt for when only provider story exists
            full_prompt = f"""
            You are a top-tier business case study writer, creating professional, detailed, and visually attractive stories for web or PDF (inspired by Storydoc, Adobe, and top SaaS companies).
 
            IMPORTANT: Write the entire case study in {detected_language}. This includes all sections, quotes, and any additional content.
 
            Only use the Solution Provider's summary below to write a complete case study. The client did not provide input. Do not label any section with "Provider Summary" or "Title". Do not include markdown (like ** or *). Just write the case study using ALL CAPS section headers and clear business English.
 
            STRUCTURE:
 
            1. LOGO & TITLE BLOCK: Display only the project title with the names of the provider and client.
            2. HERO STATEMENT / BANNER: A one-sentence summary capturing the most impactful achievement.
            3. INTRODUCTION
            4. RESEARCH AND DEVELOPMENT
            5. CLIENT CONTEXT AND CHALLENGES
            6. THE SOLUTION
            7. IMPLEMENTATION & COLLABORATION
            8. RESULTS & IMPACT
            9. CUSTOMER/CLIENT REFLECTION (create a realistic client quote based on the provider's input)
            10. TESTIMONIAL/PROVIDER REFLECTION (one quote from the provider)
            11. CALL TO ACTION
            12. QUOTES HIGHLIGHTS (2–3 extra short provider quotes NOT used earlier)
 
            CONTENT RULES:
 
            - Maintain credibility: do not fabricate specific client claims, only rephrase insights from the provider.
            - Keep each section clear and scannable using ALL CAPS headers (no bolding or markdown).
            - Include one quote in each reflection section.
            - At the end, add a "QUOTES HIGHLIGHTS" section with 2–3 additional provider quotes.
 
            Use a realistic tone and avoid generic phrases. Just output the full case study without section labels, markdown, or references to instructions.
 
            Now, here is the input:
 
            Provider Summary:
            {provider_summary}
                        

            **IMPORTANT QUOTE STRUCTURE:**
            1. **Main Story Quotes** (Only these should appear in the main story):
                - Include exactly ONE impactful quote from the provider in the "Testimonial/Provider Reflection" section
                - Create a realistic client quote for the "Customer/Client Reflection" section based on the provider's description
                - These should be the most powerful, representative quotes
                - Keep them concise and impactful

            2. **Additional Quotes** (These will appear ONLY in the meta data):
                - After the main story, provide a section titled "Quotes Highlights"
                - Include 2-3 additional meaningful quotes that were NOT used in the main story
                - These should be different from the main quotes above
                - Format each as:
                  - **Provider:** "Their exact words or close paraphrase"
                - Focus on quotes that:
                  - Highlight specific results or metrics
                  - Show unique insights about the collaboration
                  - Express satisfaction or key learnings
                  - Reveal interesting challenges overcome

            Example of Additional Quotes (for meta data only):
            - **Provider:** "The client's feedback helped us refine the solution in unexpected ways."
            """

        headers = {
            "Authorization": f"Bearer {OPENAI_API_KEY}",
            "Content-Type": "application/json"
        }

        payload = {
            "model": openai_config["model"],
            "messages": [
                {"role": "system", "content": full_prompt},
            ],
            "temperature": 0.5,
            "top_p": 0.9
        }

        response = requests.post("https://api.openai.com/v1/chat/completions", headers=headers, json=payload)
        result = response.json()
        case_study_text = result["choices"][0]["message"]["content"]
        cleaned = clean_text(case_study_text)

        def draft_quote_from_summary(summary, speaker="Client"):
            # Simple template-based fallback if OpenAI is not available
            # You can make this smarter or use OpenAI if you wish
            return f"As a {speaker.lower()}, I can say this project made a real difference for us. We're very happy with the results."

        def extract_and_remove_metadata_sections(text, client_summary=None):
            # Patterns to extract meta sections
            conflict_pattern = r"(?:\*\*|__)?Corrected\s*&\s*Conflicted Replies(?:\*\*|__)?\s*[\r\n]+(.*?)(?=(?:\*\*|__)?Quotes? Highlights(?:\*\*|__)?|$)"
            quotes_pattern = r"(?:\*\*|__)?Quotes? Highlights(?:\*\*|__)?\s*[\r\n\-:]*([\s\S]*?)(?=(?:\*\*|__)?[A-Z][^:]*:|$)"
            
            # Extract meta sections
            conflict_match = re.search(conflict_pattern, text, re.IGNORECASE | re.DOTALL)
            quotes_match = re.search(quotes_pattern, text, re.IGNORECASE | re.DOTALL)
            corrected_conflicts = conflict_match.group(1).strip() if conflict_match else ""
            quote_highlights = quotes_match.group(1).strip() if quotes_match else ""

            # Fallback: if quote_highlights is empty, try to extract blockquotes or bulleted quotes
            if not quote_highlights:
                # Try to extract lines like: - **Client:** "Quote here..."
                blockquote_lines = re.findall(r'- \*\*(Client|Provider)\*\*:\s*["""]([\s\S]*?)["""]', text)
                if blockquote_lines:
                    quote_highlights = "\n".join(f'- **{who}:** "{q.strip()}"' for who, q in blockquote_lines)
                else:
                    # Fallback: extract multi-line quotes between quotes
                    multiline_quotes = re.findall(r'["""]([\s\S]*?)["""]', text)
                    if multiline_quotes:
                        quote_highlights = "\n".join(f'- "{q.strip()}"' for q in multiline_quotes)
                    elif client_summary:
                        # Draft a quote from the client summary
                        drafted = draft_quote_from_summary(client_summary, speaker="Client")
                        quote_highlights = f'- "{drafted}"'

            # Remove meta sections from the main story
            text = re.sub(conflict_pattern, "", text, flags=re.IGNORECASE | re.DOTALL)
            text = re.sub(quotes_pattern, "", text, flags=re.IGNORECASE | re.DOTALL)
            
            # Extract key takeaways
            client_takeaways = extract_client_takeaways(client_summary) if client_summary else ""

            # Ensure sentiment analysis is included in meta data
            sentiment = analyze_sentiment(client_summary) if client_summary else {}

            return text.strip(), {
                "corrected_conflicts": corrected_conflicts,
                "quote_highlights": quote_highlights,
                "sentiment": sentiment,
                "client_takeaways": client_takeaways,
                # Add other meta data fields as needed
            }

        def extract_client_takeaways(client_summary):
            """Extract key takeaways from client interview using OpenAI."""
            try:
                prompt = f"""
                Analyze the following client interview summary and extract the 3-5 most important key takeaways.
                Focus on:
                - Main pain points or challenges they faced
                - Most valued aspects of the solution
                - Key benefits or improvements they experienced
                - Their overall satisfaction level
                - Any specific metrics or results they mentioned

                Format the response as a bullet-point list.

                Client Summary:
                {client_summary}
                """

                headers = {
                    "Authorization": f"Bearer {OPENAI_API_KEY}",
                    "Content-Type": "application/json"
                }

                payload = {
                    "model": openai_config["model"],
                    "messages": [{"role": "system", "content": prompt}],
                    "temperature": 0.3,
                    "max_tokens": 500
                }

                response = requests.post("https://api.openai.com/v1/chat/completions", 
                                      headers=headers, 
                                      json=payload)
                result = response.json()
                return result["choices"][0]["message"]["content"].strip()
            except Exception as e:
                print(f"Error extracting client takeaways: {str(e)}")
                return "Unable to extract key takeaways."

        def generate_sentiment_chart(sentiment_score, output_dir="generated_pdfs"):
            # Create a simple horizontal bar chart for sentiment
            fig, ax = plt.subplots(figsize=(4, 1.5))
            color = 'green' if sentiment_score > 6 else 'yellow' if sentiment_score > 4 else 'red'
            ax.barh(['Sentiment'], [sentiment_score], color=color)
            ax.set_xlim(0, 10)
            ax.set_xlabel('Score (0-10)')
            ax.set_title('Sentiment Score')
            plt.tight_layout()
            # Save to a unique file
            os.makedirs(output_dir, exist_ok=True)
            filename = f"sentiment_chart_{uuid.uuid4().hex}.png"
            filepath = os.path.join(output_dir, filename)
            plt.savefig(filepath)
            plt.close(fig)
            return filename

        def extract_client_satisfaction(client_summary):
            # Define categories and keywords
            categories = [
                ("Very Bad", ["terrible", "awful", "horrible", "very disappointed", "extremely dissatisfied", "never again", "worst"]),
                ("Bad", ["bad", "disappointed", "dissatisfied", "not happy", "not satisfied", "issues", "problems", "concerns"]),
                ("Neutral", ["okay", "neutral", "average", "fine", "acceptable", "neither good nor bad"]),
                ("Good", ["good", "satisfied", "happy", "pleased", "helpful", "positive", "recommend", "valuable", "improved", "great help"]),
                ("Very Good", ["excellent", "outstanding", "amazing", "fantastic", "very happy", "very satisfied", "delighted", "impressed", "exceptional", "game changer", "highly recommend", "best"])
            ]
            summary_lower = client_summary.lower()
            found_category = "Neutral"
            for cat, keywords in categories:
                for kw in keywords:
                    if kw in summary_lower:
                        found_category = cat
                        break
                if found_category != "Neutral":
                    break

            # Try to extract a satisfaction sentence
            satisfaction_sentence = ""
            for cat, keywords in categories:
                for kw in keywords:
                    match = re.search(r'([^.]*\b' + re.escape(kw) + r'\b[^.]*)\.', client_summary, re.IGNORECASE)
                    if match:
                        satisfaction_sentence = match.group(1).strip()
                        break
                if satisfaction_sentence:
                    break

            return {
                "category": found_category,
                "statement": satisfaction_sentence or "No explicit satisfaction statement found."
            }

        def generate_client_satisfaction_gauge(category):
            # Map categories to values and colors for the gauge
            category_map = {
                "Very Bad": (1, "#ef4444"),
                "Bad": (3, "#f59e42"),
                "Neutral": (5, "#fbbf24"),
                "Good": (7, "#a3e635"),
                "Very Good": (9, "#22c55e")
            }
            value, color = category_map.get(category, (5, "#fbbf24"))
            fig = go.Figure(go.Indicator(
                mode="gauge+number",
                value=value,
                number={'valueformat': '', 'font': {'size': 1}, 'suffix': ''},  # Hide the number
                title={'text': f"Client Satisfaction: <b>{category}</b>", 'font': {'size': 22}},
                gauge={
                    'axis': {'range': [0, 10], 'tickvals': [1, 3, 5, 7, 9], 'ticktext': ["Very Bad", "Bad", "Neutral", "Good", "Very Good"], 'tickwidth': 2, 'tickcolor': "#888"},
                    'bar': {'color': color, 'thickness': 0.3},
                    'steps': [
                        {'range': [0, 2], 'color': "#ef4444"},
                        {'range': [2, 4], 'color': "#f59e42"},
                        {'range': [4, 6], 'color': "#fbbf24"},
                        {'range': [6, 8], 'color': "#a3e635"},
                        {'range': [8, 10], 'color': "#22c55e"},
                    ],
                    'threshold': {
                        'line': {'color': "black", 'width': 8},
                        'thickness': 0.9,
                        'value': value
                    }
                }
            ))
            fig.update_layout(height=300, margin=dict(t=40, b=0, l=0, r=0))
            return fig.to_json()

        def analyze_sentiment(client_summary):
            try:
                analyzer = SentimentIntensityAnalyzer()
                scores = analyzer.polarity_scores(client_summary)
                compound = scores['compound']
                if compound >= 0.05:
                    sentiment = "positive"
                elif compound <= -0.05:
                    sentiment = "negative"
                else:
                    sentiment = "neutral"

                final_analysis = {
                    "overall_sentiment": {
                        "sentiment": sentiment,
                        "confidence": abs(compound),
                        "score": round((compound + 1) * 5, 2)  # scale -1..1 to 0..10
                    },
                    "emotional_analysis": {
                        "primary_emotion": sentiment,
                        "secondary_emotions": [],
                        "emotional_intensity": max(scores['pos'], scores['neg'])
                    },
                    "key_points": {
                        "positive": [],
                        "negative": []
                    },
                    "metrics": [],
                    "satisfaction": {
                        "score": 0,
                        "confidence": abs(compound),
                        "key_factors": [],
                        "statement": ""
                    },
                    "visualizations": {}
                }
                # Generate and attach the sentiment chart image
                sentiment_score = final_analysis["overall_sentiment"]["score"]
                chart_filename = generate_sentiment_chart(sentiment_score)
                final_analysis["visualizations"]["sentiment_chart_img"] = f"/generated_pdfs/{chart_filename}"

                # Add client satisfaction analysis
                satisfaction_info = extract_client_satisfaction(client_summary)
                final_analysis["satisfaction"]["category"] = satisfaction_info["category"]
                final_analysis["satisfaction"]["statement"] = satisfaction_info["statement"]

                # Generate and attach the Plotly gauge for client satisfaction
                gauge_json = generate_client_satisfaction_gauge(satisfaction_info["category"])
                final_analysis["visualizations"]["client_satisfaction_gauge"] = gauge_json

                return final_analysis
            except Exception as e:
                print(f"Error in sentiment analysis: {str(e)}")
                return {
                    "overall_sentiment": {
                        "sentiment": "unknown",
                        "confidence": 0,
                        "score": 0
                    },
                    "emotional_analysis": {
                        "primary_emotion": "unknown",
                        "secondary_emotions": [],
                        "emotional_intensity": 0
                    },
                    "key_points": {
                        "positive": [],
                        "negative": []
                    },
                    "metrics": [],
                    "satisfaction": {
                        "score": 0,
                        "confidence": 0,
                        "key_factors": [],
                        "statement": "No explicit satisfaction statement found."
                    },
                    "visualizations": {}
                }

        main_story, meta_data = extract_and_remove_metadata_sections(cleaned, client_summary)
        print("Meta data being saved:", meta_data)
        case_study.final_summary = main_story
        case_study.meta_data_text = json.dumps(meta_data, ensure_ascii=False, indent=2)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        pdf_filename = f"final_case_study_{timestamp}.pdf"
        pdf_path = os.path.join("generated_pdfs", pdf_filename)
        os.makedirs("generated_pdfs", exist_ok=True)

        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Arial", size=12)
        for line in main_story.split("\n"):
            pdf.multi_cell(0, 10, line)

        pdf.output(pdf_path)

        case_study.final_summary_pdf_path = pdf_path
        session.commit()

        return jsonify({
            "status": "success",
            "text": main_story,
            "pdf_url": f"/download/{pdf_filename}"
        })

    except Exception as e:
        session.rollback()
        return jsonify({"status": "error", "message": str(e)}), 500
    finally:
        session.close()

@app.route("/save_final_summary", methods=["POST"])
def save_final_summary():
    session = SessionLocal()
    try:
        data = request.get_json()
        case_study_id = data.get("case_study_id")
        final_summary = data.get("final_summary")

        if not case_study_id or not final_summary:
            return jsonify({"status": "error", "message": "Missing data"}), 400

        # Get the case study from DB
        case_study = session.query(CaseStudy).filter_by(id=case_study_id).first()
        if not case_study:
            return jsonify({"status": "error", "message": "Case study not found"}), 404

        # ✅ Update final summary
        case_study.final_summary = final_summary

        # ✅ Extract names from the new final summary
        names = extract_names_from_case_study(final_summary)
        lead_entity = names["lead_entity"]
        partner_entity = names["partner_entity"]
        project_title = names["project_title"]
        new_title = f"{lead_entity} x {partner_entity}: {project_title}"

        # ✅ Update CaseStudy title and name fields
        case_study.title = new_title
        case_study.provider_name = lead_entity
        case_study.client_name = partner_entity
        case_study.project_name = project_title

        session.commit()

        return jsonify({
            "status": "success",
            "message": "Final summary and title updated",
            "names": names,
            "case_study_id": case_study.id
        })

    except Exception as e:
        session.rollback()
        return jsonify({"status": "error", "message": str(e)}), 500
    finally:
        session.close()

@app.route("/download_full_summary_pdf")
def download_full_summary_pdf():
    case_study_id = request.args.get("case_study_id")
    if not case_study_id:
        return jsonify({"status": "error", "message": "Missing case_study_id"}), 400

    session = SessionLocal()
    try:
        case_study = session.query(CaseStudy).filter_by(id=case_study_id).first()

        # ✅ Check path existence
        if not case_study or not case_study.final_summary_pdf_path or not os.path.exists(case_study.final_summary_pdf_path):
            return jsonify({"status": "error", "message": "Final summary PDF not available"}), 404

        return jsonify({
            "status": "success",
            "pdf_url": f"/download/{os.path.basename(case_study.final_summary_pdf_path)}"
        })

    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500
    finally:
        session.close()

def validate_password(password):
    """Validate password strength."""
    if len(password) < 8:
        return False, "Password must be at least 8 characters long"
    if not re.search(r"[A-Z]", password):
        return False, "Password must contain at least one uppercase letter"
    if not re.search(r"[a-z]", password):
        return False, "Password must contain at least one lowercase letter"
    if not re.search(r"\d", password):
        return False, "Password must contain at least one number"
    if not re.search(r"[!@#$%^&*(),.?\":{}|<>]", password):
        return False, "Password must contain at least one special character"
    return True, ""

def validate_email(email):
    """Validate email format."""
    pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
    return bool(re.match(pattern, email))

def sanitize_input(text):
    """Sanitize user input."""
    if not text:
        return ""
    # Remove any HTML tags
    text = re.sub(r'<[^>]+>', '', text)
    # Remove any script tags
    text = re.sub(r'<script.*?>.*?</script>', '', text, flags=re.DOTALL)
    return text.strip()

@app.route('/api/signup', methods=['POST'])
def api_signup():
    data = request.get_json()
    print("DEBUG SIGNUP DATA:", data)
    required = ['first_name', 'last_name', 'email', 'company', 'password']
    if not all(data.get(f) for f in required):
        print("DEBUG: Missing field in", data)
        for f in required:
            print(f"  {f}: {data.get(f) if data else None}")
        return jsonify({'success': False, 'message': 'All fields are required.'}), 400
    session_db = SessionLocal()
    try:
        user = User(
            first_name=data['first_name'].strip(),
            last_name=data['last_name'].strip(),
            email=data['email'].strip().lower(),
            company_name=data['company'].strip(),
            password_hash=generate_password_hash(data['password'])
        )
        session_db.add(user)
        session_db.commit()
        session['user_id'] = user.id
        session.permanent = True
        return jsonify({'success': True})
    except IntegrityError:
        session_db.rollback()
        return jsonify({'success': False, 'message': 'Email already registered.'}), 409
    except Exception as e:
        session_db.rollback()
        print("DEBUG: Exception during signup:", e)
        return jsonify({'success': False, 'message': 'An error occurred during signup.'}), 500
    finally:
        session_db.close()

@app.route('/api/login', methods=['POST'])
def api_login():
    data = request.get_json()
    email = data.get('email', '').strip().lower()
    password = data.get('password', '')

    if not email or not password:
        return jsonify({'success': False, 'message': 'Email and password are required.'}), 400

    session_db = SessionLocal()
    try:
        user = session_db.query(User).filter_by(email=email).first()
        
        # Check if account is locked
        if user and user.account_locked_until and user.account_locked_until > datetime.now():
            remaining_time = (user.account_locked_until - datetime.now()).seconds // 60
            return jsonify({
                'success': False, 
                'message': f'Account is locked. Try again in {remaining_time} minutes.'
            }), 401

        if user and check_password_hash(user.password_hash, password):
            # Reset failed attempts on successful login
            user.failed_login_attempts = 0
            user.last_login = datetime.now()
            user.account_locked_until = None
            session_db.commit()
            
            session['user_id'] = user.id
            session.permanent = True
            return jsonify({'success': True})
        else:
            if user:
                # Increment failed attempts
                user.failed_login_attempts += 1
                if user.failed_login_attempts >= app.config['MAX_LOGIN_ATTEMPTS']:
                    user.account_locked_until = datetime.now() + app.config['LOGIN_LOCKOUT_DURATION']
                session_db.commit()
            
            return jsonify({'success': False, 'message': 'Invalid email or password.'}), 401
    finally:
        session_db.close()

@app.route('/api/logout', methods=['POST'])
def api_logout():
    session.clear()
    return jsonify({'success': True})

@app.route('/api/case_studies')
def api_case_studies():
    user_id = session.get('user_id')
    if not user_id:
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    label_id = request.args.get('label', type=int)
    db_session = SessionLocal()
    try:
        query = db_session.query(CaseStudy).filter_by(user_id=user_id)
        if label_id:
            query = query.join(CaseStudy.labels).filter(Label.id == label_id)
        case_studies = query.all()
        result = []
        for cs in case_studies:
            result.append({
                'id': cs.id,
                'title': cs.title,
                'solution_provider_summary': getattr(cs.solution_provider_interview, 'summary', None),
                'client_summary': getattr(cs.client_interview, 'summary', None),
                'final_summary': cs.final_summary,
                'meta_data_text': cs.meta_data_text,
                'linkedin_post': cs.linkedin_post,
                'labels': [{'id': l.id, 'name': l.name} for l in cs.labels],
                'client_link_url': getattr(cs.solution_provider_interview, 'client_link_url', None),  
                'video_url': cs.video_url,
                'video_id': cs.video_id,
                'video_status': cs.video_status,
                'video_created_at': cs.video_created_at.isoformat() if cs.video_created_at else None,
                # Pictory video fields
                'pictory_video_url': cs.pictory_video_url,
                'pictory_storyboard_id': cs.pictory_storyboard_id,
                'pictory_render_id': cs.pictory_render_id,
                'pictory_video_status': cs.pictory_video_status,
                'pictory_video_created_at': cs.pictory_video_created_at.isoformat() if cs.pictory_video_created_at else None,
                # Podcast fields
                'podcast_url': cs.podcast_url,
                'podcast_job_id': cs.podcast_job_id,
                'podcast_status': cs.podcast_status,
                'podcast_script': cs.podcast_script,
                'podcast_created_at': cs.podcast_created_at.isoformat() if cs.podcast_created_at else None,
            })
        return jsonify({'success': True, 'case_studies': result})
    finally:
        db_session.close()

@app.route('/api/labels', methods=['GET'])
def get_labels():
    user_id = session.get('user_id')
    if not user_id:
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    db_session = SessionLocal()
    try:
        labels = db_session.query(Label).filter_by(user_id=user_id).all()
        return jsonify({'success': True, 'labels': [{'id': l.id, 'name': l.name} for l in labels]})
    finally:
        db_session.close()

@app.route('/api/labels', methods=['POST'])
def create_label():
    user_id = session.get('user_id')
    if not user_id:
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    data = request.get_json()
    name = data.get('name', '').strip()
    if not name:
        return jsonify({'success': False, 'message': 'Label name required'}), 400
    db_session = SessionLocal()
    try:
        label = Label(name=name, user_id=user_id)
        db_session.add(label)
        db_session.commit()
        return jsonify({'success': True, 'label': {'id': label.id, 'name': label.name}})
    finally:
        db_session.close()

@app.route('/api/labels/<int:label_id>', methods=['PATCH'])
def rename_label(label_id):
    user_id = session.get('user_id')
    if not user_id:
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    data = request.get_json()
    new_name = data.get('name', '').strip()
    if not new_name:
        return jsonify({'success': False, 'message': 'New label name required'}), 400
    db_session = SessionLocal()
    try:
        label = db_session.query(Label).filter_by(id=label_id, user_id=user_id).first()
        if not label:
            return jsonify({'success': False, 'message': 'Label not found'}), 404
        label.name = new_name
        db_session.commit()
        return jsonify({'success': True, 'label': {'id': label.id, 'name': label.name}})
    finally:
        db_session.close()

@app.route('/api/labels/<int:label_id>', methods=['DELETE'])
def delete_label(label_id):
    user_id = session.get('user_id')
    if not user_id:
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    db_session = SessionLocal()
    try:
        label = db_session.query(Label).filter_by(id=label_id, user_id=user_id).first()
        if not label:
            return jsonify({'success': False, 'message': 'Label not found'}), 404
        db_session.delete(label)
        db_session.commit()
        return jsonify({'success': True})
    finally:
        db_session.close()

@app.route('/api/case_studies/<int:case_study_id>/labels', methods=['POST'])
def add_labels_to_case_study(case_study_id):
    user_id = session.get('user_id')
    if not user_id:
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    data = request.get_json()
    label_ids = data.get('label_ids', [])
    label_names = data.get('label_names', [])
    db_session = SessionLocal()
    try:
        case_study = db_session.query(CaseStudy).filter_by(id=case_study_id, user_id=user_id).first()
        if not case_study:
            return jsonify({'success': False, 'message': 'Case study not found'}), 404
        # Add by IDs
        for lid in label_ids:
            label = db_session.query(Label).filter_by(id=lid, user_id=user_id).first()
            if label and label not in case_study.labels:
                case_study.labels.append(label)
        # Add by names (create if not exist)
        for name in label_names:
            name = name.strip()
            if not name:
                continue
            label = db_session.query(Label).filter_by(name=name, user_id=user_id).first()
            if not label:
                label = Label(name=name, user_id=user_id)
                db_session.add(label)
                db_session.commit()
            if label not in case_study.labels:
                case_study.labels.append(label)
        db_session.commit()
        return jsonify({'success': True, 'labels': [{'id': l.id, 'name': l.name} for l in case_study.labels]})
    finally:
        db_session.close()

@app.route('/api/case_studies/<int:case_study_id>/labels/<int:label_id>', methods=['DELETE'])
def remove_label_from_case_study(case_study_id, label_id):
    user_id = session.get('user_id')
    if not user_id:
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    db_session = SessionLocal()
    try:
        case_study = db_session.query(CaseStudy).filter_by(id=case_study_id, user_id=user_id).first()
        if not case_study:
            return jsonify({'success': False, 'message': 'Case study not found'}), 404
        label = db_session.query(Label).filter_by(id=label_id, user_id=user_id).first()
        if not label or label not in case_study.labels:
            return jsonify({'success': False, 'message': 'Label not found on this case study'}), 404
        case_study.labels.remove(label)
        db_session.commit()
        return jsonify({'success': True, 'labels': [{'id': l.id, 'name': l.name} for l in case_study.labels]})
    finally:
        db_session.close()

@app.route('/api/user')
def api_user():
    user_id = session.get('user_id')
    if not user_id:
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    db_session = SessionLocal()
    try:
        user = db_session.query(User).filter_by(id=user_id).first()
        if not user:
            return jsonify({'success': False, 'message': 'User not found'}), 404
        return jsonify({
            'success': True,
            'first_name': user.first_name,
            'last_name': user.last_name,
            'email': user.email
        })
    finally:
        db_session.close()

@app.route('/api/feedback/start', methods=['POST'])
def start_feedback_session():
    user_id = session.get('user_id')
    if not user_id:
        return jsonify({'status': 'error', 'message': 'Unauthorized'}), 401
    session_id = str(uuid.uuid4())
    feedback_sessions[session_id] = {
        'user_id': user_id,
        'start_time': datetime.utcnow(),
        'transcript': [],
        'status': 'active'
    }
    return jsonify({'session_id': session_id, 'status': 'started'})

@app.route('/api/feedback/submit', methods=['POST'])
def submit_feedback():
    user_id = session.get('user_id')
    if not user_id:
        return jsonify({'status': 'error', 'message': 'Unauthorized'}), 401
    data = request.json
    session_db = SessionLocal()
    try:
        feedback = Feedback(
            user_id=user_id,
            content=data.get('content'),
            rating=data.get('rating'),
            feedback_type=data.get('feedback_type', 'general')
        )
        session_db.add(feedback)
        session_db.commit()
        return jsonify(feedback.to_dict())
    except Exception as e:
        session_db.rollback()
        return jsonify({'status': 'error', 'message': str(e)}), 500
    finally:
        session_db.close()

@app.route('/api/feedback/history', methods=['GET'])
def get_feedback_history():
    user_id = session.get('user_id')
    if not user_id:
        return jsonify({'status': 'error', 'message': 'Unauthorized'}), 401
    session_db = SessionLocal()
    try:
        feedbacks = session_db.query(Feedback).filter_by(user_id=user_id).order_by(Feedback.created_at.desc()).all()
        return jsonify([feedback.to_dict() for feedback in feedbacks])
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500
    finally:
        session_db.close()

@app.route("/get_provider_transcript", methods=["GET"])
def get_provider_transcript():
    session = SessionLocal()
    try:
        token = request.args.get("token")
        if not token:
            return jsonify({"status": "error", "message": "Missing token"}), 400

        # Get case_study_id from token
        invite = session.query(InviteToken).filter_by(token=token).first()
        if not invite:
            return jsonify({"status": "error", "message": "Invalid token"}), 404

        # Get the provider interview transcript
        provider_interview = session.query(SolutionProviderInterview).filter_by(case_study_id=invite.case_study_id).first()
        if not provider_interview or not provider_interview.transcript:
            return jsonify({"status": "error", "message": "Provider transcript not found"}), 404

        return jsonify({
            "status": "success",
            "transcript": provider_interview.transcript
        })

    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500
    finally:
        session.close()

@app.route('/generated_pdfs/<filename>')
def serve_generated_file(filename):
    return send_from_directory('generated_pdfs', filename)

def generate_linkedin_post(case_study_text):
    """Generate a LinkedIn post from a case study using AI."""
    prompt = f"""
    You are an expert LinkedIn ghostwriter for a company.  
Your job is to write **ONE well-structured LinkedIn post** that matches the exact style and tone of the sample below.

---

**🎯 Purpose:**  
The post should feel like a personal reflection from someone in the company (e.g., founder, project lead, or senior consultant) sharing a true client success story PLUS practical insights that others can learn from — written in a warm, conversational, and very clear style.

---

**✅ Write it like this:**

- Your FIRST LINE must be **one short sentence only** — maximum **10 words**.  
- It must be directly inspired by the biggest *pain point, surprising stat, or unexpected win* in the case study text.  
- It must read like a natural curiosity trigger, not a generic corporate result or claim.
• Follow immediately with 1–2 short lines that naturally set up the story or the key theme.  
• Tell a short, clear story describing:  
  — How the company worked with a client  
  — What they learned during the project  
  — How they turned those lessons into a clear, helpful framework or simple steps others can use.  
• Break down the framework or lessons just like the example: use short lines, clear steps, maybe simple arrows or numbering.  
• Optionally illustrate the *wrong way vs right way* using short lines.  
• Wrap up with 1–2 lines encouraging readers to apply the idea right away.  
• End with 3–5 relevant hashtags (all lowercase, no spaces).  
• Finally, add one line: *Visual idea:* describe a simple graphic that matches the framework.

---

**✅ Style & tone:**  
• Fully from the company's voice — "we", "our team", "our project".  
• Warm, confident, human.  
• No stiff jargon or robotic phrasing.  
• Short paragraphs, short sentences, clear line breaks — easy to read on mobile.  
• Plain language — max Grade 6–7 reading level.  
• Total length: around **1200–1800 characters**, including hashtags.

---

**❌ Do NOT:**  
• Do not use visible section labels like "HOOK".  
• Do not make it sound generic or repetitive.  
• Do not list dry bullet points — use arrows or short lines like the example.  
• Do not add any links in the post.

---

**✅ CASE STUDY:**  




    {case_study_text}

    LinkedIn Post:
    """


    headers = {
        "Authorization": f"Bearer {OPENAI_API_KEY}",
        "Content-Type": "application/json"
    }

    payload = {
        "model": "gpt-4",
        "messages": [{"role": "system", "content": prompt}],
        "temperature": 0.7,
        "max_tokens": 500
    }

    response = requests.post("https://api.openai.com/v1/chat/completions", headers=headers, json=payload)
    result = response.json()
    return result["choices"][0]["message"]["content"]

@app.route("/generate_linkedin_post", methods=["POST"])
def generate_linkedin_post_endpoint():
    session = SessionLocal()
    try:
        data = request.get_json()
        case_study_id = data.get("case_study_id")

        if not case_study_id:
            return jsonify({"status": "error", "message": "Missing case_study_id"}), 400

        case_study = session.query(CaseStudy).filter_by(id=case_study_id).first()
        if not case_study:
            return jsonify({"status": "error", "message": "Case study not found"}), 404

        if not case_study.final_summary:
            return jsonify({"status": "error", "message": "No final summary available"}), 400

        # Generate LinkedIn post
        linkedin_post = generate_linkedin_post(case_study.final_summary)
        
        # Save to database
        case_study.linkedin_post = linkedin_post
        session.commit()

        return jsonify({
            "status": "success",
            "linkedin_post": linkedin_post
        })

    except Exception as e:
        session.rollback()
        return jsonify({"status": "error", "message": str(e)}), 500
    finally:
        session.close()

def generate_heygen_input_text(final_summary):
    """Generate optimized input text for HeyGen video using OpenAI."""
    try:
        prompt = f"""You are a professional business scriptwriter creating a short video script for a HeyGen AI avatar. Your task is to turn the success story summary below into a concise, professional, and clearly structured spoken script — as if it's being presented by a company representative in a formal setting (e.g. on LinkedIn, in a client meeting, or at a company showcase).

Tone:
- Professional, confident, and factual.
- No exaggeration, slang, hype, or casual phrases (e.g., avoid words like "smashing," "amazing," "incredible," "AI-powered" unless explicitly mentioned).
- Write in first-person plural ("we," "our team," "our client") as if the company is speaking.

Content Rules:
- Only use information found in the success story summary. Do not assume or invent anything.
- Do not add filler like "since no metrics are given" — if something is missing, skip that point entirely. Every sentence must reflect real, supported content.
- Include the company name, client name, and project name where appropriate.
- Focus on the client's challenge, what was delivered, how it was implemented (if described), and the final outcome.
- If specific results or metrics are provided, include them clearly. If not, omit that section without comment.

Style:
- Use short, natural-sounding sentences — easy to follow when spoken by an AI avatar.
- The tone should sound like a real business person, not like a marketer or assistant.

Structure:
1. Brief introduction or hook (1–2 sentences)
2. Client background or challenge
3. What we delivered
4. How it was implemented (if relevant)
5. The outcome or impact
6. Closing reflection or key takeaway

Length:
- Keep the full script under 1300 characters.
- Do not include any titles, labels, line breaks, or extra notes — return only the final clean block of spoken text.

Success Story Summary:
{final_summary}

Return only the final video script. Nothing else."""

        headers = {
            "Authorization": f"Bearer {OPENAI_API_KEY}",
            "Content-Type": "application/json"
        }

        payload = {
            "model": openai_config["model"],
            "messages": [
                {"role": "system", "content": "You are a professional video script writer who specializes in creating engaging, conversational scripts for AI avatar presentations. Your scripts are known for being clear, impactful, and perfectly timed for avatar delivery."},
                {"role": "user", "content": prompt}
            ],
            "temperature": 0.7,
            "max_tokens": 500
        }

        response = requests.post("https://api.openai.com/v1/chat/completions", headers=headers, json=payload)
        result = response.json()
        script = result["choices"][0]["message"]["content"].strip()
        
        # Ensure the script doesn't exceed 1300 characters
        if len(script) > 1300:
            script = script[:1297] + "..."
            
        return script
    except Exception as e:
        print(f"Error generating HeyGen input text: {str(e)}")
        return None

def get_pictory_access_token():
    """Get access token from Pictory API."""
    try:
        headers = {
            "Content-Type": "application/json"
        }
        
        payload = {
            "client_id": PICTORY_CLIENT_ID,
            "client_secret": PICTORY_CLIENT_SECRET
        }
        
        response = requests.post(
            f"{PICTORY_API_BASE_URL}/pictoryapis/v1/oauth2/token",
            headers=headers,
            json=payload
        )
        
        if response.status_code == 200:
            return response.json().get("access_token")
        else:
            print(f"Pictory token error: {response.status_code} - {response.text}")
            return None
    except Exception as e:
        print(f"Error getting Pictory access token: {str(e)}")
        return None

def generate_pictory_scenes_text(final_summary):
    """Generate scene-based text for Pictory video using OpenAI."""
    try:
        prompt = f"""You are a video scriptwriter for StoryBoom AI. Your task is to turn the case study below into a compelling 8-scene short-form video script. Each sentence should reflect a real moment or idea from the story, written clearly enough to be visualized as a separate scene.

Your goal is to help companies showcase their project or client success story in a way that feels real, story-driven, and accurate.

Guidelines:
- Use present tense and active voice.
- Each sentence should be short (10–25 words), simple, and clear.
- Each one should reflect one key idea that can be visualized in a clip (e.g., a challenge, a solution, a result).
- Avoid vague or generic phrases like "the results were amazing." Be concrete and real.
- Use a natural, spoken tone, like someone confidently narrating their team's journey.
- Always include the company name, client name, and project name where relevant.
- Stay true to the facts and phrasing in the story. Do not exaggerate or fabricate.

Structure the 8 scenes like this:
1. A strong, curiosity-driven hook
2. The challenge or situation
3. Who we are (the solution provider)
4. What we did
5. How we delivered it
6. The main outcome
7. A highlight or metric
8. The impact on the client (or their feedback)

Output format:
Return exactly 8 sentences, separated by a period and a space. No line breaks. No bullet points. No extra text or titles.

Here is the case study:
{final_summary}

Return only the final 8-scene script, nothing else."""

        headers = {
            "Authorization": f"Bearer {OPENAI_API_KEY}",
            "Content-Type": "application/json"
        }

        payload = {
            "model": openai_config["model"],
            "messages": [
                {"role": "system", "content": "You are a professional short-form video scriptwriter who creates engaging, visual scenes for social media videos."},
                {"role": "user", "content": prompt}
            ],
            "temperature": 0.7,
            "max_tokens": 800
        }

        response = requests.post("https://api.openai.com/v1/chat/completions", headers=headers, json=payload)
        result = response.json()
        scenes_text = result["choices"][0]["message"]["content"].strip()
        
        # Split into individual scenes
        scenes = [scene.strip() for scene in scenes_text.split('\n') if scene.strip()]
        
        # Clean up numbering if present
        cleaned_scenes = []
        for scene in scenes:
            # Remove numbering like "1.", "2.", etc.
            cleaned_scene = re.sub(r'^\d+\.\s*', '', scene)
            cleaned_scenes.append(cleaned_scene)
        
        return cleaned_scenes[:6]  # Ensure max 6 scenes
    except Exception as e:
        print(f"Error generating Pictory scenes text: {str(e)}")
        return None

def create_pictory_storyboard(token, scenes, video_name):
    """Create a storyboard using Pictory API."""
    try:
        headers = {
            "Authorization": f"Bearer {token}",
            "X-Pictory-User-Id": PICTORY_USER_ID,
            "Content-Type": "application/json"
        }
        
        # Debug: Print the scenes being sent
        print(f"Generated scenes for Pictory:")
        for i, scene in enumerate(scenes, 1):
            print(f"Scene {i}: {scene}")
        
        # Create scenes array for Pictory
        # Combine all scenes into one story and let Pictory handle scene creation
        combined_story = " ".join(scenes)
        print(f"Combined story: {combined_story}")
        
        pictory_scenes = [{
            "story": combined_story,
            "createSceneOnNewLine": False,
            "createSceneOnEndOfSentence": True  # Create scenes at sentence boundaries
        }]
        
        payload = {
            "videoName": video_name,
            "videoWidth": 1080,
            "videoHeight": 1920,  # Vertical format for short-form
            "language": "en",
            "saveProject": True,
            "scenes": pictory_scenes,
            "voiceOver": {
                "enabled": True,
                "aiVoices": [
                    {
                        "speaker": "Adison",
                        "speed": 100,  # Must be >= 50 according to API
                        "amplificationLevel": 0
                    }
                ]
            },
            "backgroundMusic": {
                "enabled": True,
                "autoMusic": True,
                "volume": 0.3  # Low volume as requested
            }
        }
        
        response = requests.post(
            f"{PICTORY_API_BASE_URL}/pictoryapis/v2/video/storyboard",
            headers=headers,
            json=payload
        )
        
        if response.status_code == 200:
            return response.json().get("data", {}).get("jobId")
        else:
            print(f"Pictory storyboard error: {response.status_code} - {response.text}")
            return None
    except Exception as e:
        print(f"Error creating Pictory storyboard: {str(e)}")
        return None

def render_pictory_video(token, storyboard_job_id):
    """Render the storyboard to video using Pictory API."""
    try:
        headers = {
            "Authorization": f"Bearer {token}",
            "X-Pictory-User-Id": PICTORY_USER_ID,
            "Content-Type": "application/json"
        }
        
        response = requests.put(
            f"{PICTORY_API_BASE_URL}/pictoryapis/v2/video/render/{storyboard_job_id}",
            headers=headers
        )
        
        if response.status_code == 200:
            return response.json().get("data", {}).get("jobId")
        else:
            print(f"Pictory render error: {response.status_code} - {response.text}")
            return None
    except Exception as e:
        print(f"Error rendering Pictory video: {str(e)}")
        return None

def check_pictory_job_status(token, job_id):
    """Check the status of a Pictory job."""
    try:
        headers = {
            "Authorization": f"Bearer {token}",
            "X-Pictory-User-Id": PICTORY_USER_ID,
            "accept": "application/json",
            "Content-Type": "application/json"
        }
        
        print(f"Checking Pictory job status for job_id: {job_id}")
        print(f"Using headers: {headers}")
        
        # Use the "Get Job" endpoint from the Jobs section
        response = requests.get(
            f"{PICTORY_API_BASE_URL}/pictoryapis/v1/jobs/{job_id}",
            headers=headers
        )
        
        print(f"Pictory job status response: {response.status_code} - {response.text}")
        
        if response.status_code == 200:
            data = response.json().get("data", {})
            print(f"Pictory job data: {data}")
            return data
        else:
            print(f"Pictory job status error: {response.status_code} - {response.text}")
            return None
    except Exception as e:
        print(f"Error checking Pictory job status: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

@app.route("/api/generate_video", methods=["POST"])
def generate_video():
    session_db = SessionLocal()
    try:
        data = request.get_json()
        case_study_id = data.get('case_study_id')
        
        if not case_study_id:
            return jsonify({"error": "Case study ID is required"}), 400
            
        case_study = session_db.query(CaseStudy).filter_by(id=case_study_id).first()
        if not case_study:
            return jsonify({"error": "Case study not found"}), 404
            
        if not case_study.final_summary:
            return jsonify({"error": "Final summary is required for video generation"}), 400

        # Prevent multiple video generations for the same case study
        if case_study.video_id:
            return jsonify({"error": "A video has already been generated for this case study."}), 400

        # Generate optimized input text for HeyGen
        input_text = generate_heygen_input_text(case_study.final_summary)
        if not input_text:
            return jsonify({"error": "Failed to generate optimized input text"}), 500

        # Prepare the request to HeyGen API V2
        headers = {
            "accept": "application/json",
            "content-type": "application/json",
            "x-api-key": HEYGEN_API_KEY
        }
        
        payload = {
            "caption": False,
            "dimension": {
                "width": 1280,
                "height": 720
            },
            "video_inputs": [
                {
                    "character": {
                        "type": "avatar",
                        "avatar_id": HEYGEN_AVATAR_ID,
                        "scale": 1.0,
                        "avatar_style": "normal",
                        "offset": {
                            "x": 0.0,
                            "y": 0.0
                        }
                    },
                    "voice": {
                        "type": "text",
                        "voice_id": HEYGEN_VOICE_ID,
                        "input_text": input_text,
                        "speed": 1.0,
                        "pitch": 0
                    },
                    "background": {
                        "type": "image",
                        "url": "https://i.postimg.cc/g0tpPn1y/background3.jpg"
                    }
                }
            ]
        }

        print("Sending request to HeyGen API...")
        response = requests.post(
            f"{HEYGEN_API_BASE_URL}/video/generate",
            headers=headers,
            json=payload
        )

        print(f"HeyGen API response status: {response.status_code}")
        print(f"HeyGen API response: {response.text}")

        if response.status_code == 200:
            video_data = response.json()
            video_id = video_data.get('data', {}).get('video_id')
            
            if not video_id:
                print("No video_id in response:", video_data)
                return jsonify({
                    "status": "error",
                    "error": "No video ID received from HeyGen API"
                }), 500
            
            # Update case study with video information
            case_study.video_id = video_id
            case_study.video_status = 'processing'
            case_study.video_created_at = datetime.now(UTC)  # Use timezone-aware datetime
            session_db.commit()
            print(f"Saved video_id {video_id} to case study {case_study.id}")
            
            return jsonify({
                "status": "success",
                "video_id": video_id,
                "message": "Video generation started"
            })
        else:
            error_message = f"HeyGen API error: {response.text}"
            print(error_message)
            return jsonify({
                "status": "error",
                "error": error_message
            }), response.status_code

    except Exception as e:
        session_db.rollback()
        print(f"Error in generate_video: {str(e)}")
        return jsonify({"status": "error", "error": str(e)}), 500
    finally:
        session_db.close()

@app.route("/api/video_status/<video_id>", methods=["GET"])
def check_video_status(video_id):
    if not video_id:
        return jsonify({"error": "Video ID is required"}), 400
        
    try:
        headers = {
            "accept": "application/json",
            "x-api-key": HEYGEN_API_KEY
        }
        
        print(f"Checking status for video ID: {video_id}")
        # Use the correct v1 endpoint for status check
        response = requests.get(
            f"https://api.heygen.com/v1/video_status.get",
            headers=headers,
            params={"video_id": video_id}
        )
        
        print(f"HeyGen API response status: {response.status_code}")
        print(f"HeyGen API response: {response.text}")
        
        if response.status_code == 404:
            print("HeyGen video not ready yet (404).")
            return jsonify({"status": "not_ready", "message": "Video not ready yet"}), 200
        
        if response.status_code != 200:
            error_msg = f"HeyGen API error: {response.text}"
            print(error_msg)
            return jsonify({"error": error_msg}), 500
            
        video_data = response.json()
        print(f"HeyGen video status response: {video_data}")
        
        # Update case study with video status and URL if completed
        session_db = SessionLocal()
        try:
            # Print all video IDs in DB for debugging
            all_ids = [cs.video_id for cs in session_db.query(CaseStudy).all()]
            print("All video IDs in DB:", all_ids)
            case_study = session_db.query(CaseStudy).filter_by(video_id=video_id).first()
            
            if case_study:
                # The status is in the data object
                status = video_data.get("data", {}).get("status")
                print(f"Video status from API: {status}")
                case_study.video_status = status
                
                if status == "completed":
                    # Get the video URL from the API response
                    video_url = video_data.get("data", {}).get("video_url")
                    print(f"Video URL from API: {video_url}")
                    
                    if video_url:
                        case_study.video_url = video_url
                        session_db.commit()
                        print(f"Video URL saved to database: {case_study.video_url}")
                        return jsonify({
                            "status": "completed",
                            "video_url": video_url
                        })
                    else:
                        print("Video completed but no URL in response")
                        return jsonify({
                            "status": "completed",
                            "message": "Video completed but URL not available yet"
                        })
                elif status == "failed":
                    error = video_data.get("data", {}).get("error")
                    return jsonify({
                        "status": "failed",
                        "message": f"Video generation failed: {error}" if error else "Video generation failed"
                    })
                elif status in ["processing", "pending"]:
                    return jsonify({
                        "status": status,
                        "message": "Video is being processed"
                    })
                else:
                    # For any other status, return it as is
                    return jsonify({
                        "status": status,
                        "message": f"Video is {status}"
                    })
                
                session_db.commit()
            
            print(f"No case study found for video ID: {video_id}")
            return jsonify(video_data)
            
        except Exception as db_error:
            session_db.rollback()
            print(f"Database error: {str(db_error)}")
            return jsonify({"error": "Database error occurred"}), 500
        finally:
            session_db.close()
            
    except requests.RequestException as e:
        print(f"Request error: {str(e)}")
        return jsonify({"error": f"Failed to connect to HeyGen API: {str(e)}"}), 500
    except Exception as e:
        print(f"Unexpected error: {str(e)}")
        return jsonify({"error": f"An unexpected error occurred: {str(e)}"}), 500

@app.route("/api/generate_pictory_video", methods=["POST"])
def generate_pictory_video():
    session_db = SessionLocal()
    try:
        data = request.get_json()
        case_study_id = data.get('case_study_id')
        
        if not case_study_id:
            return jsonify({"error": "Case study ID is required"}), 400
            
        case_study = session_db.query(CaseStudy).filter_by(id=case_study_id).first()
        if not case_study:
            return jsonify({"error": "Case study not found"}), 404
            
        if not case_study.final_summary:
            return jsonify({"error": "Final summary is required for video generation"}), 400

        # Prevent multiple Pictory video generations for the same case study
        if case_study.pictory_storyboard_id:
            return jsonify({"error": "A Pictory video has already been generated for this case study."}), 400

        # Get Pictory access token
        token = get_pictory_access_token()
        if not token:
            return jsonify({"error": "Failed to get Pictory access token"}), 500

        # Generate scene-based text for Pictory
        scenes = generate_pictory_scenes_text(case_study.final_summary)
        if not scenes:
            return jsonify({"error": "Failed to generate scenes text"}), 500

        # Create video name
        video_name = f"Case Study {case_study.id} - Short Form Video"

        # Create storyboard
        storyboard_job_id = create_pictory_storyboard(token, scenes, video_name)
        if not storyboard_job_id:
            return jsonify({"error": "Failed to create Pictory storyboard"}), 500

        # Update case study with Pictory information
        case_study.pictory_storyboard_id = storyboard_job_id
        case_study.pictory_video_status = 'storyboard_processing'
        case_study.pictory_video_created_at = datetime.now(UTC)
        session_db.commit()
        
        print(f"Saved Pictory storyboard_id {storyboard_job_id} to case study {case_study.id}")
        
        return jsonify({
            "status": "success",
            "storyboard_job_id": storyboard_job_id,
            "message": "Pictory video storyboard creation started"
        })

    except Exception as e:
        session_db.rollback()
        print(f"Error in generate_pictory_video: {str(e)}")
        return jsonify({"status": "error", "error": str(e)}), 500
    finally:
        session_db.close()

@app.route("/api/pictory_video_status/<storyboard_job_id>", methods=["GET"])
def check_pictory_video_status(storyboard_job_id):
    if not storyboard_job_id:
        return jsonify({"error": "Storyboard job ID is required"}), 400
        
    try:
        print(f"Checking Pictory video status for storyboard_job_id: {storyboard_job_id}")
        
        # Get Pictory access token
        token = get_pictory_access_token()
        if not token:
            print("Failed to get Pictory access token")
            return jsonify({"error": "Failed to get Pictory access token"}), 500

        # Check storyboard status
        storyboard_status = check_pictory_job_status(token, storyboard_job_id)
        if not storyboard_status:
            print("Failed to check storyboard status")
            return jsonify({"error": "Failed to check storyboard status"}), 500

        status = storyboard_status.get("status", "unknown")
        print(f"Storyboard status: {status}")
        
        # If storyboard is completed, start rendering
        if status == "completed" and storyboard_status.get("renderParams"):
            # Get case study to check if we need to start rendering
            session_db = SessionLocal()
            try:
                case_study = session_db.query(CaseStudy).filter_by(pictory_storyboard_id=storyboard_job_id).first()
                if case_study and not case_study.pictory_render_id:
                    # Start rendering
                    render_job_id = render_pictory_video(token, storyboard_job_id)
                    if render_job_id:
                        case_study.pictory_render_id = render_job_id
                        case_study.pictory_video_status = 'rendering'
                        session_db.commit()
                        
                        return jsonify({
                            "status": "rendering",
                            "render_job_id": render_job_id,
                            "message": "Video rendering started"
                        })
                    else:
                        return jsonify({
                            "status": "error",
                            "error": "Failed to start video rendering"
                        }), 500
            finally:
                session_db.close()
        
        # Check if storyboard status already contains video URL (completed video)
        if status == "completed" and storyboard_status.get("videoURL"):
            print(f"Storyboard already contains video URL: {storyboard_status.get('videoURL')}")
            # Video is already completed in storyboard status
            session_db = SessionLocal()
            try:
                case_study = session_db.query(CaseStudy).filter_by(pictory_storyboard_id=storyboard_job_id).first()
                if case_study:
                    video_url = storyboard_status.get("videoURL")
                    case_study.pictory_video_url = video_url
                    case_study.pictory_video_status = 'completed'
                    session_db.commit()
                    
                    return jsonify({
                        "status": "completed",
                        "video_url": video_url,
                        "message": "Video is ready"
                    })
            finally:
                session_db.close()
        
        # If we have a render job, check its status
        session_db = SessionLocal()
        try:
            case_study = session_db.query(CaseStudy).filter_by(pictory_storyboard_id=storyboard_job_id).first()
            if case_study and case_study.pictory_render_id:
                print(f"Checking render job status for render_id: {case_study.pictory_render_id}")
                render_status = check_pictory_job_status(token, case_study.pictory_render_id)
                if render_status:
                    render_status_value = render_status.get("status", "unknown")
                    print(f"Render status: {render_status_value}")
                    
                    if render_status_value == "completed":
                        # Video is ready - check for video URL in various possible fields
                        video_url = (
                            render_status.get("videoURL") or  # Try videoURL first
                            render_status.get("videoUrl") or  # Try videoUrl
                            render_status.get("output", {}).get("videoUrl") or  # Try nested output.videoUrl
                            render_status.get("output", {}).get("videoURL")  # Try nested output.videoURL
                        )
                        if video_url:
                            case_study.pictory_video_url = video_url
                            case_study.pictory_video_status = 'completed'
                            session_db.commit()
                            
                            return jsonify({
                                "status": "completed",
                                "video_url": video_url,
                                "message": "Video is ready"
                            })
                        else:
                            print(f"No video URL found in render status: {render_status}")
                            return jsonify({
                                "status": "error",
                                "error": "Video completed but no URL found"
                            }), 500
                    elif render_status_value == "failed":
                        case_study.pictory_video_status = 'failed'
                        session_db.commit()
                        
                        return jsonify({
                            "status": "failed",
                            "error": "Video rendering failed"
                        }), 500
                    else:
                        return jsonify({
                            "status": "rendering",
                            "message": f"Video is {render_status_value}"
                        })
        finally:
            session_db.close()
        
        # Return storyboard status
        return jsonify({
            "status": status,
            "message": f"Storyboard is {status}"
        })
        
    except Exception as e:
        print(f"Error checking Pictory video status: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({"status": "error", "error": str(e)}), 500

def generate_podcast_prompt(final_summary):
    """Generate a podcast prompt based on the final case study summary."""
    try:
        # Extract key information from the case study
        lines = final_summary.split('\n')
        
        # Find the actual title
        title = "Business Case Study"
        for line in lines:
            if line.strip() and not line.startswith('**') and not line.startswith('-') and ':' in line:
                title = line.split(':')[0].strip()
                break
        
        # Extract a clean summary (first 800 characters)
        content = ""
        for line in lines:
            if line.strip() and not line.startswith('**') and not line.startswith('-') and len(line.strip()) > 20:
                content += line.strip() + " "
                if len(content) > 800:
                    break
        
        # Create a prompt that follows the API documentation style
        prompt = f"""Create a podcast episode titled "{title}" with an engaging, conversational tone — the kind you'd hear between two business professionals genuinely excited about a successful project. The episode should be 5–7 minutes long and sound like a natural, unscripted conversation — not a formal script or monologue.

The discussion should feel dynamic, with moments of light humor, curiosity, follow-up questions, and real reactions between the two speakers. It should reflect what it's like to listen in on a smart, energized business podcast — natural, honest, and full of small human moments.

One speaker should be the host guiding the conversation, while the other is a team member who helped deliver the project. The guest shares real stories, insights, and lessons learned from inside the solution provider's company — bringing a personal and grounded perspective to the episode.

Structure:
1. Begin with a catchy intro — a question, surprising insight, or bold hook to draw listeners in.
2. Introduce the client's background and the challenge they were facing.
3. Dive into what the team delivered — what made it effective or unique.
4. Explore how the solution was implemented — include reflections, complexity, or small wins.
5. Highlight the results and real business impact.
6. Wrap with thoughtful takeaways, lessons learned, or what they'd do differently next time.

The conversation should feel human and real — like the host and guest are riffing off each other, reacting naturally, and letting the story flow. Avoid speaker labels like "Host:" or "Guest:". Don't write a transcript or narration. Don't invent any facts.

Use only the information provided below. Return a natural, high-energy podcast episode description (max 300 words) that captures the vibe of a modern, conversational business story.

Success story summary:
{content}

Return a natural-sounding, engaging podcast episode description - it should be maximum 250 words - that captures the energy, insights, and human tone of a real business conversation.

"""



        return prompt.strip()
    except Exception as e:
        print(f"Error generating podcast prompt: {str(e)}")
        return None

@app.route("/api/generate_podcast", methods=["POST"])
def generate_podcast():
    session_db = SessionLocal()
    try:
        data = request.get_json()
        case_study_id = data.get('case_study_id')
        
        if not case_study_id:
            return jsonify({"error": "Case study ID is required"}), 400
            
        case_study = session_db.query(CaseStudy).filter_by(id=case_study_id).first()
        if not case_study:
            return jsonify({"error": "Case study not found"}), 404
            
        if not case_study.final_summary:
            return jsonify({"error": "Final summary is required for podcast generation"}), 400

        # Check if Wondercraft API key is configured
        if not WONDERCRAFT_API_KEY:
            return jsonify({"error": "Wondercraft API key not configured"}), 500

        # Clear any previous failed podcast data if this is a retry
        if case_study.podcast_status == 'failed':
            print(f"Clearing previous failed podcast data for case study {case_study.id}")
            case_study.podcast_job_id = None
            case_study.podcast_url = None
            case_study.podcast_script = None
            case_study.podcast_status = None
            case_study.podcast_created_at = None
            session_db.commit()

        # Generate podcast prompt
        podcast_prompt = generate_podcast_prompt(case_study.final_summary)
        if not podcast_prompt:
            return jsonify({"error": "Failed to generate podcast prompt"}), 500

        print(f"Generated prompt length: {len(podcast_prompt)} characters")
        print(f"Prompt preview: {podcast_prompt[:200]}...")

        # Prepare the request to Wondercraft API with correct header format
        headers = {
            "Content-Type": "application/json",
            "X-API-KEY": WONDERCRAFT_API_KEY  # Fixed header format
        }
        
        payload = {
            "prompt": podcast_prompt
        }

        print("Sending request to Wondercraft API...")
        print(f"API URL: {WONDERCRAFT_API_BASE_URL}/podcast")
        print(f"Headers: {headers}")
        
        response = requests.post(
            f"{WONDERCRAFT_API_BASE_URL}/podcast",
            headers=headers,
            json=payload,
            timeout=30
        )

        print(f"Wondercraft API response status: {response.status_code}")
        print(f"Wondercraft API response: {response.text}")

        if response.status_code == 200:
            podcast_data = response.json()
            job_id = podcast_data.get('job_id')
            
            if not job_id:
                print("No job_id in response:", podcast_data)
                return jsonify({
                    "status": "error",
                    "error": "No job ID received from Wondercraft API"
                }), 500
            
            # Update case study with podcast information
            case_study.podcast_job_id = job_id
            case_study.podcast_status = 'processing'
            case_study.podcast_created_at = datetime.now(UTC)
            session_db.commit()
            print(f"Saved podcast_job_id {job_id} to case study {case_study.id}")
            
            return jsonify({
                "status": "success",
                "job_id": job_id,
                "message": "Podcast generation started"
            })
        elif response.status_code == 429:
            return jsonify({
                "status": "error",
                "error": "Rate limit exceeded. Too many concurrent jobs. Please try again later."
            }), 429
        elif response.status_code == 422:
            return jsonify({
                "status": "error",
                "error": f"Invalid request: {response.text}"
            }), 422
        else:
            error_message = f"Wondercraft API error (Status {response.status_code}): {response.text}"
            print(error_message)
            return jsonify({
                "status": "error",
                "error": error_message
            }), response.status_code

    except Exception as e:
        session_db.rollback()
        print(f"Error in generate_podcast: {str(e)}")
        return jsonify({"status": "error", "error": str(e)}), 500
    finally:
        session_db.close()

@app.route("/api/podcast_status/<job_id>", methods=["GET"])
def check_podcast_status(job_id):
    if not job_id:
        return jsonify({"error": "Job ID is required"}), 400
        
    try:
        headers = {
            "X-API-KEY": WONDERCRAFT_API_KEY  # Fixed header format
        }
        
        print(f"Checking podcast status for job ID: {job_id}")
        response = requests.get(
            f"{WONDERCRAFT_API_BASE_URL}/podcast/{job_id}",
            headers=headers,
            timeout=10
        )
        
        print(f"Wondercraft API response status: {response.status_code}")
        print(f"Wondercraft API response: {response.text}")
        
        if response.status_code == 404:
            print("Wondercraft podcast not ready yet (404).")
            return jsonify({"status": "not_ready", "message": "Podcast not ready yet"}), 200
        
        if response.status_code != 200:
            error_msg = f"Wondercraft API error: {response.text}"
            print(error_msg)
            return jsonify({"error": error_msg}), 500
            
        podcast_data = response.json()
        print(f"Wondercraft podcast status response: {podcast_data}")
        
        # Update case study with podcast status and URL if completed
        session_db = SessionLocal()
        try:
            case_study = session_db.query(CaseStudy).filter_by(podcast_job_id=job_id).first()
            if case_study:
                status = podcast_data.get('finished', False)
                error = podcast_data.get('error', False)
                url = podcast_data.get('url')
                script = podcast_data.get('script')
                
                print(f"Status: {status}, Error: {error}, URL: {url}, Script: {script is not None}")
                
                if status and not error and url:
                    # Podcast generation completed successfully
                    case_study.podcast_status = 'completed'
                    case_study.podcast_url = url
                    case_study.podcast_script = script
                    session_db.commit()
                    print(f"Podcast completed for case study {case_study.id}")
                    
                    return jsonify({
                        "status": "completed",
                        "url": url,
                        "script": script,
                        "message": "Podcast generation completed"
                    })
                elif error:
                    # Podcast generation failed
                    case_study.podcast_status = 'failed'
                    session_db.commit()
                    print(f"Podcast generation failed for case study {case_study.id}")
                    
                    return jsonify({
                        "status": "failed",
                        "message": "Podcast generation failed",
                        "details": podcast_data
                    })
                else:
                    # Still processing
                    case_study.podcast_status = 'processing'
                    session_db.commit()
                    
                    return jsonify({
                        "status": "processing",
                        "message": "Podcast is being generated"
                    })
                
            print(f"No case study found for podcast job ID: {job_id}")
            return jsonify(podcast_data)
            
        except Exception as db_error:
            session_db.rollback()
            print(f"Database error: {str(db_error)}")
            return jsonify({"error": "Database error occurred"}), 500
        finally:
            session_db.close()
            
    except requests.RequestException as e:
        print(f"Request error: {str(e)}")
        return jsonify({"error": f"Failed to connect to Wondercraft API: {str(e)}"}), 500
    except Exception as e:
        print(f"Unexpected error: {str(e)}")
        return jsonify({"error": f"An unexpected error occurred: {str(e)}"}), 500

@app.route("/save_as_word", methods=["POST"])
def save_as_word():
    try:
        data = request.get_json()
        case_study_id = data.get("case_study_id")
        final_summary = data.get("final_summary")
        title = data.get("title", "Case Study")

        if not case_study_id or not final_summary:
            return jsonify({"status": "error", "message": "Missing case_study_id or final_summary"}), 400

        # Create Word document using python-docx
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.shared import OxmlElement, qn

        # Create a new document
        doc = Document()
        
        # Add title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.size = Inches(0.5)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add some spacing
        doc.add_paragraph()
        
        # Add the final summary content
        # Split by lines and add each line as a paragraph
        lines = final_summary.split('\n')
        for line in lines:
            if line.strip():  # Only add non-empty lines
                # Check if it's a header (all caps or starts with **)
                if line.strip().isupper() or line.strip().startswith('**'):
                    # It's a header
                    header_para = doc.add_paragraph()
                    header_run = header_para.add_run(line.strip().replace('**', ''))
                    header_run.bold = True
                    header_run.font.size = Inches(0.3)
                else:
                    # It's regular content
                    para = doc.add_paragraph()
                    para.add_run(line.strip())
        
        # Save the document to a temporary file
        import tempfile
        import os
        
        # Create a safe filename
        safe_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).rstrip()
        safe_title = safe_title.replace(' ', '_')
        filename = f"{safe_title}_{case_study_id}.docx"
        
        # Save to generated_pdfs directory (we'll use this for all generated files)
        os.makedirs("generated_pdfs", exist_ok=True)
        filepath = os.path.join("generated_pdfs", filename)
        
        doc.save(filepath)
        
        # Return the file
        return send_file(
            filepath,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        print(f"Error generating Word document: {str(e)}")
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route("/api/podcast_audio/<int:case_study_id>", methods=["OPTIONS"])
def podcast_audio_options(case_study_id):
    """Handle CORS preflight requests for podcast audio."""
    response = jsonify({"status": "ok"})
    response.headers['Access-Control-Allow-Origin'] = '*'
    response.headers['Access-Control-Allow-Methods'] = 'GET, HEAD, OPTIONS'
    response.headers['Access-Control-Allow-Headers'] = 'Range, Content-Type'
    return response

@app.route("/api/podcast_audio/<int:case_study_id>", methods=["GET"])
def serve_podcast_audio(case_study_id):
    """Proxy endpoint to serve podcast audio files to avoid CORS issues."""
    try:
        session_db = SessionLocal()
        case_study = session_db.query(CaseStudy).filter_by(id=case_study_id).first()
        
        if not case_study or not case_study.podcast_url:
            return jsonify({"error": "Podcast not found"}), 404
        
        # Fetch the audio file from the external URL
        response = requests.get(case_study.podcast_url, stream=True, timeout=30)
        
        if response.status_code != 200:
            return jsonify({"error": "Failed to fetch audio file"}), 500
        
        # Create a Flask response with the audio content
        from flask import Response
        
        def generate():
            for chunk in response.iter_content(chunk_size=8192):
                yield chunk
        
        # Return the audio as a streaming response
        return Response(
            generate(),
            content_type=response.headers.get('Content-Type', 'audio/mpeg'),
            headers={
                'Content-Length': response.headers.get('Content-Length'),
                'Accept-Ranges': 'bytes',
                'Access-Control-Allow-Origin': '*',
                'Access-Control-Allow-Methods': 'GET, HEAD, OPTIONS',
                'Access-Control-Allow-Headers': 'Range, Content-Type',
                'Cache-Control': 'public, max-age=3600'
            }
        )
        
    except Exception as e:
        print(f"Error serving podcast audio: {str(e)}")
        return jsonify({"error": "Internal server error"}), 500
    finally:
        session_db.close()

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 10000))
    app.run(host="0.0.0.0", port=port)